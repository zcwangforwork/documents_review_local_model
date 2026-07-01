"""
医疗器械体系文件审核 Agent - FastAPI 后端服务
"""
import os
import re
import json
import asyncio
import base64
import shutil
import gc
import logging
from typing import List, Dict, Optional, Any, Tuple
from contextlib import asynccontextmanager
from pathlib import Path
import tempfile

import httpx
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

# LangGraph Agent 模式导入
from agent_state import AuditState as AgentAuditState, make_initial_state
from agent_graph import (
    compile_graph as compile_agent_graph,
    compile_conversation_graph,
    create_llm as create_agent_llm,
    _classify_intent_async,
)
from langgraph.types import Command
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage

# Multi-Agent 协作模式导入（方案A · LangGraph Subagents + Send API）
from multi_agent import (
    make_multi_agent_initial_state,
    compile_supervisor_graph,
)


logger = logging.getLogger(__name__)

# ============== 配置管理 ==============
# API Key 优先从环境变量读取，不存在时使用默认值（本地开发用）
_DEFAULT_API_KEY = os.getenv("OPENAI_API_KEY", "ollama")


class Config:
    """配置管理类"""
    def __init__(self):
        # 本地 Ollama（OpenAI 兼容接口），默认运行于 http://localhost:11435
        self.api_url = os.getenv("OPENAI_API_URL", "http://localhost:11435/v1/chat/completions")
        self.embedding_url = os.getenv("EMBEDDING_API_URL", "http://localhost:11435/v1/embeddings")
        self.api_key = _DEFAULT_API_KEY
        self.model = os.getenv("OPENAI_MODEL", "qwen3.5:122b")
        self.embedding_model = os.getenv("EMBEDDING_MODEL", "qwen3-embedding:4b")
        self.timeout = float(os.getenv("REQUEST_TIMEOUT", "60"))

        # 检测是否使用了默认 API Key
        if not os.getenv("OPENAI_API_KEY"):
            print("警告: 未设置 OPENAI_API_KEY 环境变量，使用默认值。建议设置环境变量以提高安全性。")

    def to_dict(self):
        return {
            "api_url": self.api_url,
            "model": self.model,
            "timeout": self.timeout,
            "configured": bool(self.api_key)
        }


config = Config()


# ============== Pydantic 模型 ==============
class Message(BaseModel):
    """聊天消息模型"""
    role: str = Field(default="user", description="角色: user/assistant/system")
    content: str = Field(..., description="消息内容")


class ChatRequest(BaseModel):
    """聊天请求模型"""
    messages: List[Message] = Field(..., description="消息列表")
    temperature: float = Field(default=0.7, ge=0.0, le=2.0, description="温度参数")
    max_tokens: int = Field(default=32000, ge=1, le=32000, description="最大令牌数")


class ChatResponse(BaseModel):
    """聊天响应模型"""
    answer: str = Field(..., description="助手的回答")
    usage: Optional[Dict[str, Any]] = Field(None, description="令牌使用情况")


class HealthResponse(BaseModel):
    """健康检查响应"""
    status: str
    config: Dict
    vectorstore_loaded: bool = False
    document_count: int = 0


# ============== 全局变量 ==============
vector_store = None
rag_retriever = None
audit_graph = None  # LangGraph Agent 编译后的 graph (legacy mode)
conversation_graph = None  # Conversation mode graph (with SqliteSaver)
conversation_checkpointer = None  # AsyncSqliteSaver instance for conversation mode
_conversation_db_conn = None  # aiosqlite connection (closed on shutdown)
multi_agent_graph = None  # 多 Agent Supervisor Graph（方案A）
multi_agent_checkpointer = None  # 多 Agent 模式的 AsyncSqliteSaver
_multi_agent_db_conn = None  # aiosqlite connection for multi_agent mode
_session_sweeper_task: Optional[asyncio.Task] = None  # Phase 1.5 background cleanup task
# Per-session state for conversation mode: session_id -> {queue, config, graph, awaiting_input}
_conversation_sessions: Dict[str, Dict] = {}
SESSION_TIMEOUT = 7200  # 2 hours（本地大模型推理慢，放宽闲置超时）


# ============== 医疗器械 System Prompt（通用聊天用） ==============
MEDICAL_DEVICE_SYSTEM_PROMPT = """你是一个专业的贴敷式胰岛素泵生产企业文档审核数字员工（体系文件审核专家）。

你的任务是审核用户在贴敷式胰岛素泵设计开发、风险管理、软件合规、注册申报、生产质量、体系建设等方面的内部文档，结合知识库中的标准文档给出专业修改建议。

## 企业背景
用户所在企业研发和生产贴敷式胰岛素泵（patch insulin pump）— 一种可穿戴的胰岛素持续皮下输注设备，包含泵体、储液器、输注管路、嵌入式控制系统、蓝牙通信模块和配套移动端糖尿病管理APP。

## 审核范围（六大领域，覆盖设计控制全生命周期）
1. **设计开发**: ISO 13485 7.3设计控制、DHF/DMR管理，覆盖设计策划（项目计划书/可行性研究/专利分析）、设计输入（用户需求/硬件需求/结构需求/软件需求/包装需求/追溯矩阵RTM）、设计输出（硬件方案/结构方案/软件方案/编码规范/包装方案/BOM/物料规格/产品图纸/设备清单/工装图纸/检验规范/生产WI/软件版本包）、设计评审、设计验证（验证计划/性能验证/输注精度验证/包装验证/使用期限/货架有效期/运输验证/可沥滤物测试）、设计确认（临床试验/可用性测试）、设计转换（转换计划/转换报告/工艺验证计划/灭菌确认）、设计变更
2. **风险管理**: ISO 14971:2019/YY/T 0316、危害识别、初步风险分析、FMEA/DFMEA、风险控制、风险分析管理总表、网络安全风险管理、剩余风险评价、受益-风险分析、风险管理报告
3. **软件合规**: IEC 62304/YY/T 0664、软件安全分级、SDP/SRS/SADD/SDDD、软件单元/集成/系统/质量测试、软件配置管理(SCMP)、软件问题解决、软件维护、SOUP/OTS管理、网络安全测试、接口安全测试、软件追溯矩阵、网络安全追溯
4. **注册申报**: NMPA胰岛素泵注册审查指导原则、MDR 2017/745、EP清单、产品技术要求、综述资料、研究资料、临床评价报告(CER)、说明书/标签、生物相容性评价、灭菌验证、稳定性研究、第三方检测报告（生物相容性/药液相容性/安规EMC/注册检验）、供应商资质
5. **生产质量**: NMPA GMP、ISO 13485 7.5、工艺流程图、工艺验证计划/方案/报告(IQ/OQ/PQ)、批生产记录(BMR)、来料/过程/成品检验规范、设备管理、供应商管理、灭菌批记录/灭菌确认/灭菌工艺验证、工装验收、生产检验SOP、UDI标识管理
6. **体系建设**: ISO 13485:2016全体系、质量手册、程序文件、作业指导书、记录表单、设计控制程序、风险管理程序、软件开发程序、文件控制、CAPA、内审、管理评审、培训记录、PMS上市后监督

## 审核原则
1. **完整性检查**：文件是否覆盖相关法规条款的全部要求
2. **一致性检查**：文件内容是否相互协调、无矛盾
3. **可操作性**：文件描述是否足够具体、可执行
4. **证据链**：是否有相应的记录表单支撑执行证据
5. **产品特异性**：是否充分考虑了贴敷式胰岛素泵的特殊风险和控制措施

## 重要提示
- 直接给出审核结果，不要输出思考过程
- 回答应专业、具体、可操作
- 建议用户上传文件选择对应的专项审核领域以获得最精准的审核结果
- 可审核的六大领域：设计开发、风险管理、软件合规、注册申报、生产质量、体系建设
"""


# ============== 会话管理 ==============
class ConversationHistory:
    """会话历史管理（带总大小限制，防止多轮对话累积导致内存溢出）"""
    def __init__(self, max_history: int = 10):
        self.history: Dict[str, List[Dict]] = {}
        self.max_history = max_history
        self.max_content_length = 8000   # 单条消息最大字符数
        self.max_total_chars = 50000     # 单个session总字符数上限

    def get_or_create(self, session_id: str = "default") -> List[Dict]:
        if session_id not in self.history:
            self.history[session_id] = []
        return self.history[session_id]

    def add_message(self, session_id: str, role: str, content: str):
        messages = self.get_or_create(session_id)
        if len(content) > self.max_content_length:
            content = content[:self.max_content_length] + f"\n... [已截断，原长度:{len(content)}字符]"
        messages.append({"role": role, "content": content})

        # 按消息数量裁剪
        if len(messages) > self.max_history:
            system_msg = [m for m in messages if m["role"] == "system"]
            other_msgs = [m for m in messages if m["role"] != "system"]
            self.history[session_id] = system_msg + other_msgs[-self.max_history:]

        # 按总字符数裁剪：超出上限时从旧消息开始删除
        total_chars = sum(len(m["content"]) for m in self.history[session_id])
        if total_chars > self.max_total_chars:
            trimmed = []
            running = 0
            for m in reversed(self.history[session_id]):
                running += len(m["content"])
                if m["role"] == "system" or running <= self.max_total_chars:
                    trimmed.insert(0, m)
            self.history[session_id] = trimmed

    def clear(self, session_id: str = "default"):
        if session_id in self.history:
            del self.history[session_id]


conversation_manager = ConversationHistory()


# ============== 逐段审核状态管理 ==============
class SegmentState:
    """单个会话的逐段审核状态"""
    def __init__(self, document_text: str, filename: str, audit_type: str, doc_type: str, segment_size: int = 4000):
        self.document_text = document_text
        self.filename = filename
        self.audit_type = audit_type
        self.doc_type = doc_type
        self.segment_size = segment_size
        self.current_position = 0
        self.segment_results: List[Dict] = []
        self.total_segments = self._calculate_total_segments()

    def _calculate_total_segments(self) -> int:
        """计算文档将被分成多少段"""
        if not self.document_text:
            return 0
        total = 0
        pos = 0
        text = self.document_text
        while pos < len(text):
            end = self._find_segment_end(text, pos)
            total += 1
            pos = end
        return total

    def _find_segment_end(self, text: str, start: int) -> int:
        """找到段落结束位置，尽量在段落或句子边界断开"""
        target = min(start + self.segment_size, len(text))
        if target >= len(text):
            return len(text)

        # 在 target 附近查找最佳断开点
        search_end = min(target + 500, len(text))
        # 优先级：双换行 > 单换行 > 句号 > target 位置
        best = target
        search_start = max(start + self.segment_size // 2, start)

        # 1. 双换行（段落边界）
        pos = text.find('\n\n', search_start, search_end)
        if pos != -1 and pos < target + 300:
            return pos + 2

        # 2. 单换行
        pos = text.find('\n', search_start, search_end)
        if pos != -1 and pos < target + 200:
            return pos + 1

        # 3. 句号后
        for p in ['. ', '。', '！', '？', '! ', '? ']:
            pos = text.rfind(p, search_start, search_end)
            if pos != -1 and abs(pos - target) < 300:
                return pos + len(p)

        return target

    def get_next_segment(self) -> Optional[Tuple[str, int, int, int, int]]:
        """
        获取下一个文本段落

        Returns:
            (text, segment_index, total_segments, start_pos, end_pos) 或 None（已完成）
        """
        if self.current_position >= len(self.document_text):
            return None

        start = self.current_position
        end = self._find_segment_end(self.document_text, start)
        segment_text = self.document_text[start:end].strip()
        self.current_position = end

        segment_index = len(self.segment_results) + 1
        return (segment_text, segment_index, self.total_segments, start, end)

    def add_result(self, result: Dict):
        """保存段落审核结果"""
        self.segment_results.append(result)

    def get_progress(self) -> Dict:
        """获取当前进度信息"""
        total_chars = len(self.document_text)
        progress_pct = round(self.current_position / total_chars * 100, 1) if total_chars > 0 else 100.0
        return {
            "current_position": self.current_position,
            "total_chars": total_chars,
            "progress_pct": progress_pct,
            "segments_completed": len(self.segment_results),
            "total_segments": self.total_segments,
            "is_complete": self.current_position >= total_chars
        }


class SegmentManager:
    """逐段审核状态管理器"""
    def __init__(self):
        self._states: Dict[str, SegmentState] = {}

    def create(self, session_id: str, document_text: str, filename: str, audit_type: str, doc_type: str, segment_size: int = 4000) -> SegmentState:
        state = SegmentState(document_text, filename, audit_type, doc_type, segment_size)
        self._states[session_id] = state
        return state

    def get(self, session_id: str) -> Optional[SegmentState]:
        return self._states.get(session_id)

    def remove(self, session_id: str):
        if session_id in self._states:
            del self._states[session_id]
            gc.collect()


segment_manager = SegmentManager()


def smart_split_text(text: str, segment_size: int = 4000) -> List[Tuple[int, int]]:
    """
    将文本按 segment_size 智能分割，返回每段的 (start, end) 位置列表。
    尽量在段落边界或句子边界断开。
    """
    segments = []
    pos = 0
    while pos < len(text):
        target = min(pos + segment_size, len(text))
        if target >= len(text):
            segments.append((pos, len(text)))
            break

        search_end = min(target + 500, len(text))
        search_start = max(pos + segment_size // 2, pos)
        best = target

        # 1. 双换行（段落边界）
        found = text.find('\n\n', search_start, search_end)
        if found != -1 and found < target + 300:
            best = found + 2
        else:
            # 2. 单换行
            found = text.find('\n', search_start, search_end)
            if found != -1 and found < target + 200:
                best = found + 1
            else:
                # 3. 句号后
                for p in ['. ', '。', '！', '？', '! ', '? ']:
                    found = text.rfind(p, search_start, search_end)
                    if found != -1 and abs(found - target) < 300:
                        best = found + len(p)
                        break

        segments.append((pos, best))
        pos = best

    return segments


# ============== 初始化向量存储 ==============
def init_vector_store():
    """初始化向量存储和 RAG 检索器（带内存监控）"""
    global vector_store, rag_retriever

    try:
        from vector_store import create_vector_store, MiniMaxEmbeddingFunction
        from rag_retriever import create_rag_retriever

        base_dir = os.path.dirname(os.path.abspath(__file__))
        db_path = os.path.join(base_dir, "data", "chroma_db_insulin_pump")

        # 检查数据库大小并告警
        db_file = os.path.join(db_path, "chroma.sqlite3")
        if os.path.exists(db_file):
            db_size_mb = os.path.getsize(db_file) / (1024 * 1024)
            print(f"向量库文件大小: {db_size_mb:.0f} MB")
            if db_size_mb > 1000:
                print(f"[WARNING] 向量库文件超过1GB ({db_size_mb:.0f} MB)，但查询已配置为仅使用v2 collection，不影响运行")

        # 检查系统可用内存
        try:
            import psutil
            avail_mb = psutil.virtual_memory().available / (1024 * 1024)
            print(f"系统可用内存: {avail_mb:.0f} MB")
            if avail_mb < 2048:
                print(f"[WARNING] 可用内存不足2GB，可能导致服务不稳定")
        except ImportError:
            pass

        embedding_function = MiniMaxEmbeddingFunction(
            api_key=config.api_key,
            api_url=config.embedding_url,
            model=config.embedding_model,
            dimension=1024
        )

        vector_store = create_vector_store(persist_directory=db_path, embedding_function=embedding_function)
        rag_retriever = create_rag_retriever(
            vector_store=vector_store,
            api_key=config.api_key,
            api_url=config.api_url,
            model=config.model
        )
        return True
    except Exception as e:
        print(f"向量存储初始化失败: {e}")
        import traceback
        traceback.print_exc()
        return False


# ============== API 调用函数 ==============
async def call_openai_api(messages: List[Dict], temperature: float = 0.7, max_tokens: int = 32000) -> Dict:
    """调用 OpenAI 兼容 API，带重试机制"""
    if not config.api_key:
        raise HTTPException(status_code=500, detail="API Key 未配置")

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config.api_key}"
    }

    payload = {
        "model": config.model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens
    }

    last_error = None
    max_retries = 3

    async with httpx.AsyncClient(timeout=config.timeout, trust_env=False) as client:
        for attempt in range(max_retries):
            try:
                response = await client.post(config.api_url, headers=headers, json=payload)
                response.raise_for_status()
                # 安全解析 JSON：API 可能返回非 JSON 错误文本（如 "Internal Server Error"）
                try:
                    return response.json()
                except json.JSONDecodeError:
                    response_text = response.text[:500]
                    print(f"[API] 非 JSON 响应 (尝试 {attempt+1}/{max_retries}): {response_text}")
                    if attempt < max_retries - 1:
                        wait_time = 2 ** attempt
                        print(f"[API] 等待 {wait_time} 秒后重试...")
                        await asyncio.sleep(wait_time)
                        continue
                    raise HTTPException(
                        status_code=502,
                        detail=f"API 返回了非 JSON 格式的响应: {response_text}"
                    )
            except httpx.HTTPStatusError as e:
                if e.response.status_code == 429 and attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"[API] 速率限制 (尝试 {attempt+1}/{max_retries}), 等待 {wait_time} 秒...")
                    await asyncio.sleep(wait_time)
                    continue
                # 安全读取错误响应体
                try:
                    error_detail = e.response.json() if e.response.content else {"error": str(e)}
                except json.JSONDecodeError:
                    error_detail = {"error": e.response.text[:500] if e.response.content else str(e)}
                raise HTTPException(status_code=e.response.status_code, detail=error_detail)
            except httpx.TimeoutException as e:
                last_error = e
                if attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"[API] 请求超时 (尝试 {attempt+1}/{max_retries}), 等待 {wait_time} 秒...")
                    await asyncio.sleep(wait_time)
                    continue
            except httpx.ConnectError as e:
                last_error = e
                if attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"[API] 连接失败 (尝试 {attempt+1}/{max_retries}), 等待 {wait_time} 秒...")
                    await asyncio.sleep(wait_time)
                    continue
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"API 调用失败: {str(e)}")

        if last_error:
            raise HTTPException(status_code=504, detail=f"API 请求多次失败: {str(last_error)}")
        raise HTTPException(status_code=502, detail="API 请求失败，已重试多次")


async def get_embeddings(texts: List[str]) -> List[List[float]]:
    """获取文本嵌入向量"""
    if not config.api_key:
        raise HTTPException(status_code=500, detail="API Key 未配置")

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config.api_key}"
    }

    embeddings = []
    async with httpx.AsyncClient(timeout=60, trust_env=False) as client:
        for text in texts:
            payload = {
                "model": config.embedding_model,
                "input": text[:8000]
            }
            try:
                response = await client.post(config.embedding_url, headers=headers, json=payload)
                if response.status_code == 200:
                    result = response.json()
                    data = result.get("data", [])
                    if data:
                        embedding = data[0].get("embedding", [])
                        # 截断到 1024 维（与向量库维度保持一致）
                        if len(embedding) > 1024:
                            embedding = embedding[:1024]
                        embeddings.append(embedding)
                    else:
                        embeddings.append([])
                else:
                    embeddings.append([])
            except Exception:
                embeddings.append([])

    return embeddings


# ============== FastAPI 应用 ==============
@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期管理"""
    print("贴敷式胰岛素泵文档审核数字员工 后端服务启动")
    print(f"API URL: {config.api_url}")
    print(f"Model: {config.model}")

    # 初始化向量存储
    vs_loaded = init_vector_store()
    if vs_loaded and vector_store:
        doc_count = vector_store.count()
        print(f"向量库已加载: {doc_count} 个文档")
    else:
        print("警告: 向量库未加载，部分功能可能不可用")

    # 初始化 LangGraph Agent (legacy one-shot mode)
    global audit_graph
    try:
        agent_llm = create_agent_llm()
        audit_graph = compile_agent_graph(llm=agent_llm, retriever=rag_retriever)
        print("LangGraph Agent (legacy) 已编译就绪")
    except Exception as e:
        print(f"警告: LangGraph Agent 初始化失败: {e}")
        audit_graph = None

    # 初始化 Conversation Graph (对话模式, 带 AsyncSqliteSaver)
    global conversation_graph, conversation_checkpointer, _conversation_db_conn
    try:
        import aiosqlite
        from langgraph.checkpoint.sqlite.aio import AsyncSqliteSaver

        os.makedirs("checkpoints", exist_ok=True)
        _conversation_db_conn = await aiosqlite.connect("checkpoints/conversation.db")
        conversation_checkpointer = AsyncSqliteSaver(_conversation_db_conn)
        conv_llm = create_agent_llm()
        conversation_graph = compile_conversation_graph(
            llm=conv_llm, retriever=rag_retriever,
            checkpointer=conversation_checkpointer,
        )
        print("LangGraph Conversation Graph 已编译就绪 (AsyncSqliteSaver)")
    except Exception as e:
        print(f"警告: Conversation Graph 初始化失败: {e}")
        import traceback
        traceback.print_exc()
        conversation_graph = None
        conversation_checkpointer = None

    # 初始化 Multi-Agent Supervisor Graph（方案A · 章节级并行 Send API）
    global multi_agent_graph, multi_agent_checkpointer, _multi_agent_db_conn
    try:
        import aiosqlite
        from langgraph.checkpoint.sqlite.aio import AsyncSqliteSaver

        os.makedirs("checkpoints", exist_ok=True)
        _multi_agent_db_conn = await aiosqlite.connect("checkpoints/multi_agent.db")
        multi_agent_checkpointer = AsyncSqliteSaver(_multi_agent_db_conn)
        ma_llm = create_agent_llm()
        multi_agent_graph = compile_supervisor_graph(
            llm=ma_llm,
            retriever=rag_retriever,
            checkpointer=multi_agent_checkpointer,
        )
        print("LangGraph Multi-Agent Supervisor Graph 已编译就绪 (AsyncSqliteSaver)")
    except Exception as e:
        print(f"警告: Multi-Agent Graph 初始化失败: {e}")
        import traceback
        traceback.print_exc()
        multi_agent_graph = None
        multi_agent_checkpointer = None

    # Phase 1.5: 启动 session 清理任务
    global _session_sweeper_task
    _session_sweeper_task = asyncio.create_task(_session_cleanup_loop())
    print(f"Session 清理任务已启动 (timeout={SESSION_TIMEOUT}s, interval={SESSION_SWEEP_INTERVAL}s)")

    yield

    # 关闭共享的 httpx 客户端
    if rag_retriever:
        await rag_retriever.close()

    # 取消 session 清理后台任务
    if _session_sweeper_task and not _session_sweeper_task.done():
        _session_sweeper_task.cancel()
        try:
            await _session_sweeper_task
        except (asyncio.CancelledError, Exception):
            pass

    # 关闭对话模式的 aiosqlite 连接
    if _conversation_db_conn:
        await _conversation_db_conn.close()
        print("Conversation DB 连接已关闭")

    # 关闭多 Agent 模式的 aiosqlite 连接
    if _multi_agent_db_conn:
        await _multi_agent_db_conn.close()
        print("Multi-Agent DB 连接已关闭")

    print("贴敷式胰岛素泵文档审核数字员工 后端服务关闭")


app = FastAPI(
    title="贴敷式胰岛素泵文档审核数字员工 API",
    description="贴敷式胰岛素泵企业文档智能审核系统 — 六大领域（设计开发/风险管理/软件合规/注册申报/生产质量/体系建设）",
    version="3.1.0",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FRONTEND_DIR = os.path.join(os.path.dirname(BASE_DIR), "frontend")


# ============== 前端路由 ==============
@app.get("/")
async def serve_frontend():
    """服务前端页面"""
    return FileResponse(os.path.join(FRONTEND_DIR, "index.html"))


@app.get("/agent")
async def serve_agent_frontend():
    """服务 Agent 审核专用页面"""
    agent_html = os.path.join(FRONTEND_DIR, "agent.html")
    if os.path.exists(agent_html):
        return FileResponse(agent_html)
    return FileResponse(os.path.join(FRONTEND_DIR, "index.html"))


app.mount("/static", StaticFiles(directory=FRONTEND_DIR), name="static")


# ============== API 路由 ==============
@app.get("/info", response_model=Dict)
async def root():
    """根路径 - 服务信息"""
    return {
        "name": "贴敷式胰岛素泵文档审核数字员工 API",
        "version": "3.1.0",
        "docs": "/docs",
        "audit_domains": [
            "risk_management", "design_dev", "software_compliance",
            "registration", "production_quality", "system_construction"
        ]
    }


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """健康检查接口"""
    doc_count = vector_store.count() if vector_store else 0
    return HealthResponse(
        status="healthy",
        config=config.to_dict(),
        vectorstore_loaded=vector_store is not None,
        document_count=doc_count
    )


@app.post("/api/chat", response_model=ChatResponse)
async def chat(request: ChatRequest, session_id: str = "default"):
    """聊天接口 - 通用问答模式"""
    all_messages = [{"role": "system", "content": MEDICAL_DEVICE_SYSTEM_PROMPT}]

    history = conversation_manager.get_or_create(session_id)
    all_messages.extend(history)

    for msg in request.messages:
        all_messages.append({"role": msg.role, "content": msg.content})

    result = await call_openai_api(
        messages=all_messages,
        temperature=request.temperature,
        max_tokens=request.max_tokens
    )

    # OpenAI 兼容格式响应解析
    answer = ""
    choices = result.get("choices", [])
    if choices and len(choices) > 0:
        choice = choices[0]
        message = choice.get("message", {})
        answer = message.get("content", "")
        # DeepSeek-V4-Pro 模型：当 content 为空时，从 reasoning_content 提取实际回答
        if not answer:
            reasoning = message.get("reasoning_content", "")
            if reasoning:
                # 将思考过程作为回答返回（虽不完美，但优于空响应）
                answer = reasoning

    if not answer:
        answer = str(result) if result else ""

    for msg in request.messages:
        conversation_manager.add_message(session_id, msg.role, msg.content)
    conversation_manager.add_message(session_id, "assistant", answer)

    return ChatResponse(answer=answer, usage=result.get("usage"))


@app.post("/api/upload")
async def upload_document(
    file: UploadFile = File(...),
    session_id: str = Form("default")
):
    """
    上传体系文件并提取文本内容

    Args:
        file: 上传的文件（.docx 或 .pdf）
        session_id: 会话 ID

    Returns:
        提取的文档内容和基本信息
    """
    # 检查文件类型
    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ['.docx', '.pdf']:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}，仅支持 .docx 和 .pdf"
        )

    # 保存上传文件到临时目录（流式写入，避免全量加载到内存）
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            # 分块流式写入，每块64KB，避免大文件OOM
            while True:
                chunk = await file.read(64 * 1024)  # 64KB chunks
                if not chunk:
                    break
                tmp.write(chunk)
            tmp_path = tmp.name
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    try:
        # 提取文本（结构化 Markdown 格式）
        from doc_processor import extract_text, get_file_metadata

        text = extract_text(tmp_path)
        metadata = get_file_metadata(tmp_path)
        metadata["filename"] = filename

        if not text or len(text.strip()) < 50:
            raise HTTPException(status_code=400, detail="文档内容过少或无法提取文本")

        # 保存到会话历史（截断防止会话历史 OOM）
        max_history_text = min(len(text), 5000)
        conversation_manager.add_message(session_id, "user", f"[上传文件: {filename}]\n{text[:max_history_text]}")
        conversation_manager.add_message(session_id, "assistant", f"已收到文件: {filename}，文档长度: {len(text)} 字符。请问您想如何处理这个文件？")

        return {
            "filename": filename,
            "text_length": len(text),
            "text_preview": text[:1000],
            "metadata": metadata,
            "session_id": session_id
        }

    finally:
        # 删除临时文件
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


@app.post("/api/analyze")
async def analyze_document(
    file: UploadFile = File(...),
    question: str = Form("请审核这份体系文件，给出修改建议"),
    session_id: str = Form("default"),
    audit_type: str = Form("risk_management"),
    doc_type: str = Form("")
):
    """
    使用多轮审核流水线分析体系文件

    流水线：章节分割 → 逐章并发审核 → 综合分析

    Args:
        file: 上传的文件（.docx 或 .pdf）
        question: 用户的问题或指令
        session_id: 会话 ID
        audit_type: 审核类型，"risk_management"（风险管理专项）或 "general"（综合体系审核）

    Returns:
        审核结果
    """
    if not rag_retriever:
        raise HTTPException(status_code=503, detail="向量库未加载，请稍后重试")

    # 日志：开始处理时的内存状态
    try:
        import psutil
        proc = psutil.Process(os.getpid())
        start_mem_mb = proc.memory_info().rss / (1024 * 1024)
        print(f"[analyze] 开始处理，当前进程内存: {start_mem_mb:.0f} MB")
    except ImportError:
        pass

    # 检查文件类型
    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ['.docx', '.pdf']:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    # 验证审核类型（六大领域 + general 保底）
    VALID_AUDIT_TYPES = [
        "risk_management", "design_dev", "software_compliance",
        "registration", "production_quality", "system_construction", "general"
    ]
    if audit_type not in VALID_AUDIT_TYPES:
        audit_type = "general"

    # 保存上传文件到临时目录（流式写入，避免全量加载到内存）
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            while True:
                chunk = await file.read(64 * 1024)  # 64KB chunks
                if not chunk:
                    break
                tmp.write(chunk)
            tmp_path = tmp.name
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    try:
        from doc_processor import extract_text

        # 提取文档文本（结构化 Markdown 格式）
        text = extract_text(tmp_path)
        if not text:
            raise HTTPException(status_code=400, detail="无法提取文档文本")

        # 使用多轮审核流水线分析文档
        result = await rag_retriever.analyze_document(
            user_document=text,
            user_filename=filename,
            audit_type=audit_type,
            doc_type=doc_type
        )

        # 构建检索到的文档信息
        retrieved_docs_info = []
        seen_sources = set()
        for doc in result.get("retrieved_docs", []):
            source = doc.get("source", "未知")
            if source not in seen_sources:
                seen_sources.add(source)
                retrieved_docs_info.append({
                    "source": source,
                    "preview": doc.get("text", "")[:200]
                })

        # 保存到会话历史（截断防止 OOM）
        max_answer_len = min(len(result["answer"]), 5000)
        conversation_manager.add_message(session_id, "user", f"[上传文件分析: {filename}]")
        conversation_manager.add_message(session_id, "assistant", result["answer"][:max_answer_len])

        # 日志：处理完成时的内存状态
        try:
            import psutil
            proc = psutil.Process(os.getpid())
            end_mem_mb = proc.memory_info().rss / (1024 * 1024)
            print(f"[analyze] 处理完成，当前进程内存: {end_mem_mb:.0f} MB (增加 {end_mem_mb - start_mem_mb:.0f} MB)")
        except Exception:
            pass

        return {
            "filename": filename,
            "answer": result["answer"],
            "section_count": result.get("section_count", 0),
            "section_results": result.get("section_results", []),
            "usage": None,
            "retrieved_docs": retrieved_docs_info,
            "audit_type": audit_type,
            "pipeline": "multi_pass"
        }

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ============== Agent 审核 API ==============
@app.post("/api/agent/analyze")
async def agent_analyze(
    file: UploadFile = File(...),
    session_id: str = Form("default"),
    audit_type: str = Form("risk_management"),
    doc_type: str = Form("")
):
    """
    Agent 模式审核 — 同步端点，FastAPI 自动在线程池中执行
    """
    if not audit_graph:
        raise HTTPException(status_code=503, detail="Agent 审核服务未就绪，请检查 LangGraph 初始化")

    if not rag_retriever:
        raise HTTPException(status_code=503, detail="向量库未加载，请稍后重试")

    VALID_AUDIT_TYPES = [
        "risk_management", "design_dev", "software_compliance",
        "registration", "production_quality", "system_construction", "general"
    ]
    if audit_type not in VALID_AUDIT_TYPES:
        audit_type = "general"

    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ['.docx', '.pdf']:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            contents = file.file.read()
            tmp.write(contents)
            tmp_path = tmp.name
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    try:
        from doc_processor import extract_text, parse_document_structure

        text = extract_text(tmp_path)
        if not text or len(text.strip()) < 50:
            raise HTTPException(status_code=400, detail="文档内容过少或无法提取文本")

        # 使用 LLM 解析文档结构（正则回退）
        # request_timeout=600: 本地 qwen3.5:122b（thinking 默认开启）解析大文档单次推理可能需数分钟；doc_processor 内置重试
        outline, sections = parse_document_structure(text, llm=create_agent_llm(temperature=0.1, request_timeout=600, max_tokens=16384, streaming=False))

        initial_state = make_initial_state(
            document_text=text,
            document_type=audit_type,
            outline=outline,
            sections=sections,
        )

        # recursion_limit = 25(底数) + 章节数×5(每节audit+evaluate+update+reaudit)
        config = {
            "configurable": {"thread_id": session_id},
            "recursion_limit": max(200, 25 + len(sections) * 5),
        }
        print(f"[Agent] 开始审核: {filename}, 审核类型: {audit_type}, 章节数: {len(sections)}, recursion_limit: {config['recursion_limit']}")

        final_state = await audit_graph.ainvoke(initial_state, config)

        # 提取结果
        final_report = final_state.get("final_report", "")
        audit_results = final_state.get("audit_results", [])
        contradictions = final_state.get("contradictions", [])

        # 构建检索文档信息
        retrieved_docs_info = []
        if audit_results:
            seen_sources = set()
            for r in audit_results:
                findings = r.get("findings", [])
                for f in findings[:3]:
                    if f not in seen_sources:
                        seen_sources.add(f)
                        retrieved_docs_info.append({"source": f, "preview": f[:200]})

        print(f"[Agent] 审核完成: {len(audit_results)} 个小节, 报告长度: {len(final_report)} 字符")

        return {
            "filename": filename,
            "answer": final_report,
            "section_count": len(sections),
            "section_results": audit_results,
            "audit_type": audit_type,
            "pipeline": "agent",
            "contradictions": contradictions,
            "retrieved_docs": retrieved_docs_info,
        }

    except Exception as e:
        logger.error(f"Agent 审核失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Agent 审核失败: {str(e)}")

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ============== 对话式审核 API (Conversation Mode) ==============

class ConversationSession:
    """对话式审核会话状态"""
    def __init__(self, session_id: str, config: dict):
        self.session_id = session_id
        self.config = config
        self.queue: asyncio.Queue = asyncio.Queue()
        self.awaiting_input = False
        self.task: Optional[asyncio.Task] = None
        self.filename = ""
        self.resume_event: asyncio.Event = asyncio.Event()
        self.pending_message: str = ""  # User message to resume with
        self.queued_message: str = ""  # Message sent while graph is running (picked up at next checkpoint)
        # Phase 1.5: timestamps for the session-timeout sweeper.
        import time as _time
        now = _time.time()
        self.created_at: float = now
        self.last_active: float = now

    def touch(self) -> None:
        """Mark this session as recently used (called on any user/SSE activity)."""
        import time as _time
        self.last_active = _time.time()


async def _cleanup_session(session_id: str):
    """清理会话资源"""
    if session_id in _conversation_sessions:
        session = _conversation_sessions[session_id]
        if session.task and not session.task.done():
            session.task.cancel()
        _conversation_sessions.pop(session_id, None)


# Phase 1.5: background sweeper for stale conversation sessions.
# Scans every SESSION_SWEEP_INTERVAL seconds and removes any session whose
# `last_active` exceeds SESSION_TIMEOUT — preventing memory growth from
# abandoned uploads.
SESSION_SWEEP_INTERVAL = int(os.getenv("SESSION_SWEEP_INTERVAL", "120"))


async def _session_cleanup_loop():
    import time as _time
    while True:
        try:
            await asyncio.sleep(SESSION_SWEEP_INTERVAL)
            now = _time.time()
            expired = [
                sid for sid, s in list(_conversation_sessions.items())
                if (now - getattr(s, "last_active", now)) > SESSION_TIMEOUT
            ]
            for sid in expired:
                logger.info(f"[session-sweeper] cleaning up expired session: {sid}")
                await _cleanup_session(sid)
        except asyncio.CancelledError:
            raise
        except Exception as e:
            logger.warning(f"[session-sweeper] loop iteration failed: {e}")


# Auto-resume timeout (seconds) for progress checkpoints.
# The user can send a message during this window to interrupt.
AUTO_RESUME_TIMEOUT = 3.0


async def _run_graph_and_push_events(
    session: ConversationSession,
    initial_input,
):
    """
    Persistent background task: runs the conversation graph continuously.

    - **progress** checkpoints: auto-resume after AUTO_RESUME_TIMEOUT seconds
      unless the user sends a message during the window.
    - **strategy_review / post_audit** checkpoints: wait indefinitely for user input.
    - SSE stays connected across all checkpoints.
    """
    graph = conversation_graph
    state_to_run = initial_input

    try:
        while True:
            # ----- run graph until interrupt or completion -----
            # stream_mode=["messages","updates"] yields tuples (mode, payload).
            # - "messages": LLM token chunks → push as `token` events for live UI.
            # - "updates": node-completion deltas → push as `node_complete` events.
            async for mode, chunk in graph.astream(
                state_to_run, session.config,
                stream_mode=["messages", "updates"],
            ):
                if mode == "messages":
                    message_chunk, metadata = chunk
                    token_text = getattr(message_chunk, "content", "") or ""
                    if token_text:
                        await session.queue.put({
                            "event": "token",
                            "content": token_text,
                            "node": (metadata or {}).get("langgraph_node", ""),
                        })
                    # Phase 3.4: Tool 调用事件 — Agent 触发工具时上报给前端
                    tool_calls = getattr(message_chunk, "tool_calls", None) or []
                    for tc in tool_calls:
                        if isinstance(tc, dict):
                            await session.queue.put({
                                "event": "tool_call",
                                "tool_name": tc.get("name", ""),
                                "tool_args": tc.get("args", {}),
                                "tool_id": tc.get("id", ""),
                                "node": (metadata or {}).get("langgraph_node", ""),
                            })
                    # ToolMessage：工具调用结果
                    msg_type = type(message_chunk).__name__
                    if msg_type == "ToolMessage":
                        tool_output = getattr(message_chunk, "content", "") or ""
                        if isinstance(tool_output, list):
                            tool_output = " ".join(str(c) for c in tool_output)
                        await session.queue.put({
                            "event": "tool_result",
                            "tool_name": getattr(message_chunk, "name", ""),
                            "tool_id": getattr(message_chunk, "tool_call_id", ""),
                            "result_preview": str(tool_output)[:500],
                            "node": (metadata or {}).get("langgraph_node", ""),
                        })
                elif mode == "updates":
                    for node_name, node_output in chunk.items():
                        current_stage = (
                            node_output.get("current_stage", "")
                            if isinstance(node_output, dict) else ""
                        )
                        await session.queue.put({
                            "event": "node_complete",
                            "node": node_name,
                            "current_stage": current_stage,
                        })

            # ----- check if graph was interrupted -----
            current = await graph.aget_state(session.config)
            if current and current.interrupts:
                interrupt_values = []
                for interrupt_item in current.interrupts:
                    interrupt_values.append(interrupt_item.value)

                # Use the last interrupt value
                value = interrupt_values[-1] if interrupt_values else {}
                checkpoint_type = value.get("checkpoint", "")

                await session.queue.put({
                    "event": "checkpoint",
                    "checkpoint": checkpoint_type,
                    "stage": value.get("stage", ""),
                    "summary": value.get("summary", ""),
                    "section_count": value.get("section_count", 0),
                    "audited_count": value.get("audited_count", 0),
                    "completed": value.get("completed", 0),
                    "total": value.get("total", 0),
                    "structured": value.get("structured", {}),
                })

                session.awaiting_input = True

                if checkpoint_type == "progress":
                    # ---- Auto-resume: wait briefly for user interruption ----
                    await session.queue.put({
                        "event": "awaiting_input",
                        "checkpoint": checkpoint_type,
                        "auto_resume": True,
                        "auto_resume_timeout": AUTO_RESUME_TIMEOUT,
                    })

                    try:
                        await asyncio.wait_for(
                            session.resume_event.wait(),
                            timeout=AUTO_RESUME_TIMEOUT,
                        )
                        # User sent a message — use it
                        message = session.pending_message or "继续"
                    except asyncio.TimeoutError:
                        # No user message — auto-resume
                        message = "继续"
                        # Also check for queued messages from earlier
                        if session.queued_message:
                            message = session.queued_message
                            session.queued_message = ""
                            await session.queue.put({
                                "event": "agent_message",
                                "content": f"已收到你的消息，Agent 正在处理...",
                            })

                    session.resume_event.clear()
                    session.pending_message = ""
                    session.awaiting_input = False
                    await session.queue.put({"event": "resuming"})
                    state_to_run = Command(resume=message)
                    continue

                else:
                    # ---- Major checkpoint: wait indefinitely for user ----
                    await session.queue.put({
                        "event": "awaiting_input",
                        "checkpoint": checkpoint_type,
                        "auto_resume": False,
                    })

                    # Merge any queued message into pending context
                    if session.queued_message:
                        if not session.pending_message:
                            session.pending_message = session.queued_message
                        session.queued_message = ""

                    await session.resume_event.wait()
                    session.resume_event.clear()

                    message = session.pending_message or "继续"
                    session.pending_message = ""
                    session.awaiting_input = False
                    state_to_run = Command(resume=message)
                    continue

            # ----- no interrupt — graph completed -----
            final_state = await graph.aget_state(session.config)
            report = ""
            if final_state and final_state.values:
                report = final_state.values.get("final_report", "")
            await session.queue.put({
                "event": "complete",
                "report": report,
            })
            session.awaiting_input = False
            break

    except asyncio.CancelledError:
        pass
    except Exception as e:
        logger.error(f"Graph execution error [{session.session_id}]: {e}")
        import traceback
        traceback.print_exc()
        await session.queue.put({"event": "error", "message": str(e)})


def _build_agent_summary(audit_plan: dict, sections: list, audit_results: list = None) -> str:
    """Build a summary message shown to user at strategy checkpoint"""
    parts = [f"文档解析完成，共 **{len(sections)}** 个章节。\n"]
    deep = sum(1 for v in audit_plan.values() if v.get("depth") == "deep")
    standard = sum(1 for v in audit_plan.values() if v.get("depth") == "standard")
    quick = sum(1 for v in audit_plan.values() if v.get("depth") == "quick")
    parts.append(f"审核策略: deep={deep}, standard={standard}, quick={quick}\n")
    if audit_results:
        total_findings = sum(len(r.get("findings", [])) for r in audit_results)
        parts.append(f"审核完成，共发现 **{total_findings}** 个问题项。")
    parts.append("\n请回复您的意见：继续 / 重审指定章节 / 跳过章节 / 补充信息")
    return "".join(parts)


@app.post("/api/agent/conversation/start")
async def conversation_start(
    file: Optional[UploadFile] = File(None),
    session_id: str = Form("default"),
    audit_type: str = Form("risk_management"),
    doc_type: str = Form(""),
):
    """
    启动对话式审核

    上传文档 → 解析结构 → 创建 conversation session → 返回 session_id
    前端随后连接 SSE (/stream/{session_id}) 接收实时事件

    支持无文件启动：未上传文档时进入自由对话模式（不启动 graph，
    awaiting_input=True，用户消息直接由 LLM 回答）。
    """
    if not conversation_graph:
        raise HTTPException(status_code=503, detail="对话式审核服务未就绪")

    if not rag_retriever:
        raise HTTPException(status_code=503, detail="向量库未加载，请稍后重试")

    VALID_TYPES = [
        "risk_management", "design_dev", "software_compliance",
        "registration", "production_quality", "system_construction", "general"
    ]
    if audit_type not in VALID_TYPES:
        audit_type = "general"

    # ===== 自由对话模式（无文件）=====
    if file is None or not file.filename:
        await _cleanup_session(session_id)
        config = {
            "configurable": {"thread_id": session_id},
            "recursion_limit": 100,
        }
        session = ConversationSession(session_id=session_id, config=config)
        session.filename = ""
        session.awaiting_input = True  # 直接进入待输入状态，等待用户消息
        _conversation_sessions[session_id] = session

        # 推送欢迎 checkpoint 事件，前端会展示提示并启用输入
        await session.queue.put({
            "event": "checkpoint",
            "checkpoint": "free_chat",
            "stage": "free_chat",
            "summary": "已进入自由对话模式（未上传文档）。你可以直接提问医疗器械法规、体系文件、审核标准等问题，我会基于专业知识回答。如需审核文档，请刷新页面重新开始并上传文件。",
            "structured": {},
        })

        print(f"[Conversation] 自由对话会话已创建: {session_id}")
        return {
            "session_id": session_id,
            "filename": "",
            "section_count": 0,
            "audit_type": audit_type,
            "mode": "free_chat",
            "message": "自由对话会话已创建，请通过 SSE 连接接收事件",
        }

    # ===== 文档审核模式（原逻辑）=====
    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ['.docx', '.pdf']:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    # Clean up any existing session with same ID
    await _cleanup_session(session_id)

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            contents = file.file.read()
            tmp.write(contents)
            tmp_path = tmp.name
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    try:
        from doc_processor import extract_text, parse_document_structure

        text = extract_text(tmp_path)
        if not text or len(text.strip()) < 50:
            raise HTTPException(status_code=400, detail="文档内容过少或无法提取文本")

        # 使用 LLM 解析文档结构（正则回退）
        # request_timeout=600: 本地 qwen3.5:122b（thinking 默认开启）解析大文档单次推理可能需数分钟；doc_processor 内置重试
        outline, sections = parse_document_structure(text, llm=create_agent_llm(temperature=0.1, request_timeout=600, max_tokens=16384, streaming=False))

        initial_state = make_initial_state(
            document_text=text,
            document_type=audit_type,
            outline=outline,
            sections=sections,
            conversation_mode=True,
        )

        config = {
            "configurable": {"thread_id": session_id},
            "recursion_limit": max(200, 25 + len(sections) * 5),
        }

        session = ConversationSession(session_id=session_id, config=config)
        session.filename = filename
        _conversation_sessions[session_id] = session

        # Start graph execution in background
        session.task = asyncio.create_task(
            _run_graph_and_push_events(session, initial_state)
        )

        print(f"[Conversation] 会话已创建: {session_id}, 章节数: {len(sections)}")

        return {
            "session_id": session_id,
            "filename": filename,
            "section_count": len(sections),
            "audit_type": audit_type,
            "message": "会话已创建，请通过 SSE 连接接收实时事件",
        }

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


@app.get("/api/agent/conversation/stream/{session_id}")
async def conversation_stream(session_id: str):
    """
    SSE 端点：推送 Agent 实时状态

    事件类型:
    - node_complete: 节点执行完成 {event, node, current_stage}
    - checkpoint: Agent 等待用户反馈 {event, checkpoint, stage, summary}
    - complete: 审核完成 {event, report}
    - error: 执行错误 {event, message}
    """

    if session_id not in _conversation_sessions:
        # SSE error events don't have status codes, send error event
        async def error_stream():
            data = json.dumps({"event": "error", "message": "会话不存在或已过期"}, ensure_ascii=False)
            yield f"data: {data}\n\n"
        return StreamingResponse(error_stream(), media_type="text/event-stream")

    session = _conversation_sessions[session_id]
    session.touch()

    async def event_stream():
        try:
            while True:
                try:
                    event = await asyncio.wait_for(session.queue.get(), timeout=60.0)
                    session.touch()
                    yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"

                    # Only break on terminal events — checkpoint is NOT terminal
                    if event.get("event") in ("complete", "error"):
                        break
                except asyncio.TimeoutError:
                    # 后台审核长时间无事件（如本地大模型推理慢）时，
                    # SSE 仍保持连接：心跳同时刷新 last_active，避免被
                    # session-sweeper 误判为过期而清理，导致"会话不存在或已过期"
                    session.touch()
                    yield f"data: {json.dumps({'event': 'heartbeat'})}\n\n"

        except asyncio.CancelledError:
            pass

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/agent/conversation/message")
async def conversation_message(
    session_id: str = Form(...),
    message: str = Form(...),
):
    """
    Send a message to the agent — works at ANY time, even during graph execution.

    - If the graph is at a checkpoint: process the message and resume immediately.
    - If the graph is running: queue the message; it will be picked up at the
      next checkpoint (progress checkpoints happen every ~5 sections).
    """
    if session_id not in _conversation_sessions:
        raise HTTPException(status_code=404, detail="会话不存在或已过期")

    session = _conversation_sessions[session_id]
    session.touch()

    if not conversation_graph:
        raise HTTPException(status_code=503, detail="对话式审核服务未就绪")

    # Step 1: Classify user intent
    intent_llm = create_agent_llm(temperature=0.1)
    classification = await _classify_intent_async(message, intent_llm)
    intent = classification.get("intent", "approve")

    logger.info(
        f"[Conversation] {session_id} intent={intent}, awaiting_input={session.awaiting_input}, "
        f"re_audit_sections={classification.get('re_audit_sections', [])}"
    )

    # ===== 自由对话模式（无文档）：强制 question intent，直接 LLM 回答 =====
    # 自由对话模式下 graph 未启动，跳过所有 graph 状态更新与 resume 逻辑。
    if session.awaiting_input and not session.filename:
        question_text = (
            classification.get("question_text", message)
            if classification.get("intent") == "question" else message
        )
        try:
            answer_llm = create_agent_llm(temperature=0.3)
            answer_prompt = (
                "你是一个医疗器械文档审核专家，精通 ISO 13485、IEC 62304、ISO 14971、"
                "MDR 2017/745、NMPA GMP 等法规标准。\n\n"
                f"用户问题：{question_text}\n\n"
                "请基于专业知识详细回答。如果用户的问题需要审核具体文档，"
                "请提示用户上传文档后使用审核功能。"
            )
            answer_resp = await answer_llm.ainvoke([HumanMessage(content=answer_prompt)])
            extra_response = (
                answer_resp.content if hasattr(answer_resp, "content") else str(answer_resp)
            )
        except Exception as e:
            logger.warning(f"Free-chat answering failed: {e}")
            extra_response = f"抱歉，处理您的问题时出错: {e}"

        await session.queue.put({
            "event": "agent_message",
            "content": extra_response,
            "intent": "question",
        })
        return {
            "session_id": session_id,
            "intent": "question",
            "status": "answered",
            "message": "消息已处理",
            "extra_response": extra_response,
        }

    if not session.awaiting_input:
        # Graph is running — queue the message for the next checkpoint
        session.queued_message = message
        await session.queue.put({
            "event": "agent_message",
            "content": f"已收到你的消息：「{message[:100]}」\nAgent 将在当前批次审核完成后立即处理。",
        })
        return {
            "session_id": session_id,
            "intent": intent,
            "status": "queued",
            "message": "消息已排队，将在下一个检查点处理",
        }

    # Graph is at a checkpoint — handle the message now
    extra_response = None

    if intent == "question":
        # Answer the user's question directly, then auto-resume the graph
        question_text = classification.get("question_text", message)
        try:
            answer_llm = create_agent_llm(temperature=0.3)
            answer_prompt = (
                f"你是一个医疗器械文档审核专家。用户在当前审核过程中问了一个问题。\n\n"
                f"用户问题：{question_text}\n\n"
                f"请简洁回答（100字以内）。"
            )
            answer_resp = await answer_llm.ainvoke([HumanMessage(content=answer_prompt)])
            extra_response = answer_resp.content if hasattr(answer_resp, "content") else str(answer_resp)
            await session.queue.put({
                "event": "agent_message",
                "content": extra_response,
                "intent": "question",
            })
        except Exception as e:
            logger.warning(f"Question answering failed: {e}")
            await session.queue.put({
                "event": "agent_message",
                "content": f"抱歉，处理您的问题时出错: {e}",
                "intent": "question",
            })

    elif intent == "adjust":
        re_audit_sections = classification.get("re_audit_sections", [])
        standard_override = classification.get("standard_override", "")
        revision_requests = []
        for idx in re_audit_sections:
            revision_requests.append({
                "section_idx": idx,
                "request": message,
                "standard_override": standard_override,
            })
        await conversation_graph.aupdate_state(
            session.config,
            {
                "re_audit_sections": re_audit_sections,
                "revision_requests": revision_requests,
                "user_feedback": message,
            },
        )

    elif intent == "supplement":
        supplement_text = classification.get("supplement_text", message)
        current = await conversation_graph.aget_state(session.config)
        existing_context = ""
        if current and current.values:
            existing_context = current.values.get("conversation_context", "")
        new_context = (
            f"{existing_context}\n[User Feedback]: {supplement_text}"
            if existing_context
            else f"[User Feedback]: {supplement_text}"
        )
        await conversation_graph.aupdate_state(
            session.config,
            {
                "conversation_context": new_context[:8000],
                "user_feedback": message,
            },
        )

    elif intent == "skip":
        skip_sections = classification.get("skip_sections", [])
        await conversation_graph.aupdate_state(
            session.config,
            {
                "skip_sections": skip_sections,
                "user_feedback": message,
            },
        )

    elif intent == "regenerate":
        instruction = classification.get("regenerate_instruction", message)
        await conversation_graph.aupdate_state(
            session.config,
            {
                "user_feedback": message,
                "conversation_context": f"[Regenerate instruction]: {instruction}",
            },
        )

    # Step 3: Signal the persistent graph loop to resume
    session.pending_message = message
    session.resume_event.set()

    return {
        "session_id": session_id,
        "intent": intent,
        "classification": classification,
        "message": "消息已处理，Agent 正在继续执行",
        "extra_response": extra_response,
    }


@app.post("/api/agent/conversation/reaudit")
async def conversation_reaudit(
    session_id: str = Form(...),
    sections: str = Form(...),  # Comma-separated section indices e.g. "2,5,7"
    standard_override: str = Form(""),
):
    """
    触发指定章节的重新审核

    这是一个便捷端点，等同于发送 "重审第X节" 消息
    """
    if session_id not in _conversation_sessions:
        raise HTTPException(status_code=404, detail="会话不存在或已过期")

    session = _conversation_sessions[session_id]
    session.touch()
    if not conversation_graph:
        raise HTTPException(status_code=503, detail="对话式审核服务未就绪")

    try:
        section_indices = [int(s.strip()) for s in sections.split(",") if s.strip()]
    except ValueError:
        raise HTTPException(status_code=400, detail="sections 格式错误，应为逗号分隔数字")

    if not section_indices:
        raise HTTPException(status_code=400, detail="请指定要重新审核的章节索引")

    # Check re_audit cycle limit
    current_state = await conversation_graph.aget_state(session.config)
    cycle_count = 0
    if current_state and current_state.values:
        cycle_count = current_state.values.get("re_audit_cycle_count", 0)
    if cycle_count >= 3:
        raise HTTPException(status_code=400, detail=f"已达到最大重审次数 ({cycle_count}/3)")

    revision_requests = [{
        "section_idx": idx,
        "request": f"Re-audit section {idx}",
        "standard_override": standard_override,
    } for idx in section_indices]

    await conversation_graph.aupdate_state(
        session.config,
        {
            "re_audit_sections": section_indices,
            "revision_requests": revision_requests,
            "user_feedback": f"重审章节: {sections}",
        },
    )

    # Signal the persistent graph loop to resume
    session.pending_message = f"重审章节: {sections}"
    session.resume_event.set()

    return {
        "session_id": session_id,
        "re_audit_sections": section_indices,
        "standard_override": standard_override,
        "cycle_count": cycle_count + 1,
        "message": f"正在重新审核 {len(section_indices)} 个章节",
    }


@app.get("/api/agent/conversation/state/{session_id}")
async def conversation_state(session_id: str):
    """查询当前对话式审核的状态"""
    if session_id not in _conversation_sessions:
        raise HTTPException(status_code=404, detail="会话不存在或已过期")

    session = _conversation_sessions[session_id]
    if not conversation_graph:
        raise HTTPException(status_code=503, detail="对话式审核服务未就绪")

    current = await conversation_graph.aget_state(session.config)
    state_values = current.values if current else {}

    return {
        "session_id": session_id,
        "filename": session.filename,
        "awaiting_input": session.awaiting_input,
        "current_stage": state_values.get("current_stage", ""),
        "pending_checkpoint": state_values.get("pending_checkpoint", ""),
        "section_count": len(state_values.get("sections", [])),
        "audited_count": len(state_values.get("audit_results", [])),
        "re_audit_cycle_count": state_values.get("re_audit_cycle_count", 0),
        "finished": state_values.get("finished", False),
        "final_report": state_values.get("final_report", ""),
        "has_interrupts": bool(current.interrupts if current else False),
    }


@app.get("/api/agent/conversation/report/{session_id}")
async def conversation_report(session_id: str):
    """获取对话式审核的最终报告"""
    if session_id not in _conversation_sessions:
        raise HTTPException(status_code=404, detail="会话不存在或已过期")

    session = _conversation_sessions[session_id]
    if not conversation_graph:
        raise HTTPException(status_code=503, detail="对话式审核服务未就绪")

    current = await conversation_graph.aget_state(session.config)
    state_values = current.values if current else {}

    report = state_values.get("final_report", "")
    if not report:
        raise HTTPException(status_code=404, detail="报告尚未生成")

    return {
        "session_id": session_id,
        "report": report,
        "final_report": report,
        "finished": state_values.get("finished", False),
    }


# ============== Multi-Agent 协作模式 API（方案A）==============
# Structure Analyzer → Chapter Auditor ×N (Send API fan-out) → Report Synthesizer

# Per-session state for multi-agent conversation mode
_multi_agent_sessions: Dict[str, "MultiAgentSession"] = {}


class MultiAgentSession:
    """多 Agent 协作模式的对话会话状态。

    与单 Agent ConversationSession 类似，但状态字段对齐 MultiAgentState。
    多 Agent 当前不在 supervisor 图内插入 interrupt 检查点，仅做 SSE 流式推送
    （node_complete / chapter_progress / token / complete / error）。
    """

    def __init__(self, session_id: str, config: dict):
        self.session_id = session_id
        self.config = config
        self.queue: asyncio.Queue = asyncio.Queue()
        self.task: Optional[asyncio.Task] = None
        self.filename = ""
        import time as _time
        now = _time.time()
        self.created_at: float = now
        self.last_active: float = now

    def touch(self) -> None:
        import time as _time
        self.last_active = _time.time()


async def _cleanup_multi_agent_session(session_id: str):
    if session_id in _multi_agent_sessions:
        session = _multi_agent_sessions[session_id]
        if session.task and not session.task.done():
            session.task.cancel()
        _multi_agent_sessions.pop(session_id, None)


@app.post("/api/multi-agent/analyze")
async def multi_agent_analyze(
    file: UploadFile = File(...),
    session_id: str = Form("default"),
    audit_type: str = Form("risk_management"),
    doc_type: str = Form(""),
):
    """多 Agent 协作模式 — 一次性同步审核（章节级并行）。

    流程: 上传 → 解析结构 → analyze_structure_agent → 章节级并行 Chapter Auditor
          → cross_validate → synthesize_report → 返回 final_report
    """
    if not multi_agent_graph:
        raise HTTPException(
            status_code=503,
            detail="Multi-Agent 协作服务未就绪，请检查 supervisor_graph 初始化"
        )

    if not rag_retriever:
        raise HTTPException(status_code=503, detail="向量库未加载，请稍后重试")

    VALID_AUDIT_TYPES = [
        "risk_management", "design_dev", "software_compliance",
        "registration", "production_quality", "system_construction", "general"
    ]
    if audit_type not in VALID_AUDIT_TYPES:
        audit_type = "general"

    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ['.docx', '.pdf']:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            contents = file.file.read()
            tmp.write(contents)
            tmp_path = tmp.name
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    try:
        from doc_processor import extract_text, parse_document_structure

        text = extract_text(tmp_path)
        if not text or len(text.strip()) < 50:
            raise HTTPException(status_code=400, detail="文档内容过少或无法提取文本")

        # 使用 LLM 解析文档结构（正则回退）
        outline, sections = parse_document_structure(
            text,
            llm=create_agent_llm(temperature=0.1, request_timeout=600, max_tokens=16384, streaming=False)
        )

        initial_state = make_multi_agent_initial_state(
            document_text=text,
            document_type=audit_type,
            outline=outline,
            sections=sections,
        )

        # recursion_limit: 多 Agent 模式只有 3 层节点（structure/audit/synthesize）+ Send fan-out，
        # 但保守起见仍按章节数线性放大
        config = {
            "configurable": {"thread_id": f"multi-agent-{session_id}"},
            "recursion_limit": max(100, 25 + len(sections) * 3),
        }

        print(
            f"[MultiAgent] 开始审核: {filename}, 审核类型: {audit_type}, "
            f"章节数: {len(sections)}, recursion_limit: {config['recursion_limit']}"
        )

        final_state = await multi_agent_graph.ainvoke(initial_state, config)

        final_report = final_state.get("final_report", "")
        chapter_results = final_state.get("chapter_results", [])
        subsection_results = final_state.get("subsection_results", [])
        contradictions = final_state.get("contradictions", [])
        outline_summary = final_state.get("outline_summary", {})

        print(
            f"[MultiAgent] 审核完成: 章节数={len(chapter_results)}, "
            f"小节数={len(subsection_results)}, 报告长度={len(final_report)}"
        )

        return {
            "filename": filename,
            "answer": final_report,
            "section_count": len(sections),
            "chapter_count": len(chapter_results),
            "section_results": subsection_results,
            "chapter_results": [
                {
                    "chapter_idx": c.get("chapter_idx"),
                    "chapter_title": c.get("chapter_title"),
                    "section_count": len(c.get("subsection_results", [])),
                    "chapter_summary": c.get("chapter_summary", ""),
                }
                for c in chapter_results
            ],
            "outline_summary": outline_summary,
            "contradictions": contradictions,
            "audit_type": audit_type,
            "pipeline": "multi_agent",
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Multi-Agent 审核失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Multi-Agent 审核失败: {str(e)}")

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


async def _run_multi_agent_and_push_events(
    session: MultiAgentSession,
    initial_input: dict,
):
    """多 Agent SSE 后台任务：流式推送 supervisor graph 执行进度。

    事件类型:
    - token        : LLM 流式 token（章节审核 / 综合报告）
    - node_complete: supervisor 节点完成（analyze_structure / audit_chapter / synthesize_report）
    - chapter_done : 单个章节审核完成（携带 chapter_idx, chapter_title, section_count）
    - complete     : 全部完成，携带 final_report
    - error        : 异常
    """
    graph = multi_agent_graph
    try:
        async for mode, chunk in graph.astream(
            initial_input, session.config,
            stream_mode=["messages", "updates"],
        ):
            if mode == "messages":
                message_chunk, metadata = chunk
                token_text = getattr(message_chunk, "content", "") or ""
                if token_text:
                    await session.queue.put({
                        "event": "token",
                        "content": token_text,
                        "node": (metadata or {}).get("langgraph_node", ""),
                    })
            elif mode == "updates":
                for node_name, node_output in chunk.items():
                    if not isinstance(node_output, dict):
                        continue

                    # 章节审核完成（reducer 视角下每次会单独 push 一份 chapter_results 增量）
                    if node_name == "audit_chapter":
                        chapters_delta = node_output.get("chapter_results", []) or []
                        for ch in chapters_delta:
                            await session.queue.put({
                                "event": "chapter_done",
                                "chapter_idx": ch.get("chapter_idx"),
                                "chapter_title": ch.get("chapter_title"),
                                "section_count": len(ch.get("subsection_results", [])),
                                "chapter_summary": (ch.get("chapter_summary", "") or "")[:500],
                            })
                    else:
                        await session.queue.put({
                            "event": "node_complete",
                            "node": node_name,
                            "current_stage": node_output.get("current_stage", ""),
                            "total_chapters": node_output.get("total_chapters"),
                            "total_sections": node_output.get("total_sections"),
                            "total_agent_steps": node_output.get("total_agent_steps"),
                        })

        # graph 执行完成
        final_state = await graph.aget_state(session.config)
        report = ""
        chapter_count = 0
        section_count = 0
        if final_state and final_state.values:
            report = final_state.values.get("final_report", "") or ""
            chapter_count = len(final_state.values.get("chapter_results", []) or [])
            section_count = len(final_state.values.get("subsection_results", []) or [])
        await session.queue.put({
            "event": "complete",
            "report": report,
            "chapter_count": chapter_count,
            "section_count": section_count,
        })

    except asyncio.CancelledError:
        pass
    except Exception as e:
        logger.error(f"Multi-Agent graph execution error [{session.session_id}]: {e}")
        import traceback
        traceback.print_exc()
        await session.queue.put({"event": "error", "message": str(e)})


@app.post("/api/multi-agent/conversation/start")
async def multi_agent_conversation_start(
    file: UploadFile = File(...),
    session_id: str = Form("default"),
    audit_type: str = Form("risk_management"),
    doc_type: str = Form(""),
):
    """启动多 Agent 协作流式审核会话。

    上传文档 → 解析结构 → 创建 multi-agent session → 启动后台 graph 执行
    前端连接 SSE (/api/multi-agent/conversation/stream/{session_id}) 接收事件
    """
    if not multi_agent_graph:
        raise HTTPException(status_code=503, detail="Multi-Agent 协作服务未就绪")

    if not rag_retriever:
        raise HTTPException(status_code=503, detail="向量库未加载，请稍后重试")

    VALID_TYPES = [
        "risk_management", "design_dev", "software_compliance",
        "registration", "production_quality", "system_construction", "general"
    ]
    if audit_type not in VALID_TYPES:
        audit_type = "general"

    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ['.docx', '.pdf']:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    # 清理可能存在的旧会话
    await _cleanup_multi_agent_session(session_id)

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            contents = file.file.read()
            tmp.write(contents)
            tmp_path = tmp.name
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    try:
        from doc_processor import extract_text, parse_document_structure

        text = extract_text(tmp_path)
        if not text or len(text.strip()) < 50:
            raise HTTPException(status_code=400, detail="文档内容过少或无法提取文本")

        outline, sections = parse_document_structure(
            text,
            llm=create_agent_llm(temperature=0.1, request_timeout=600, max_tokens=16384, streaming=False)
        )

        initial_state = make_multi_agent_initial_state(
            document_text=text,
            document_type=audit_type,
            outline=outline,
            sections=sections,
            conversation_mode=True,
        )

        config = {
            "configurable": {"thread_id": f"multi-agent-{session_id}"},
            "recursion_limit": max(100, 25 + len(sections) * 3),
        }

        session = MultiAgentSession(session_id=session_id, config=config)
        session.filename = filename
        _multi_agent_sessions[session_id] = session

        session.task = asyncio.create_task(
            _run_multi_agent_and_push_events(session, initial_state)
        )

        print(
            f"[MultiAgent:Conversation] 会话已创建: {session_id}, "
            f"章节数(待 supervisor 重新分章): sections={len(sections)}"
        )

        return {
            "session_id": session_id,
            "filename": filename,
            "section_count": len(sections),
            "audit_type": audit_type,
            "pipeline": "multi_agent",
            "message": "会话已创建，请通过 SSE 连接接收实时事件",
        }

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


@app.get("/api/multi-agent/conversation/stream/{session_id}")
async def multi_agent_conversation_stream(session_id: str):
    """多 Agent 协作模式 SSE 端点。

    事件: token / node_complete / chapter_done / complete / error / heartbeat
    """
    if session_id not in _multi_agent_sessions:
        async def error_stream():
            data = json.dumps({"event": "error", "message": "会话不存在或已过期"}, ensure_ascii=False)
            yield f"data: {data}\n\n"
        return StreamingResponse(error_stream(), media_type="text/event-stream")

    session = _multi_agent_sessions[session_id]
    session.touch()

    async def event_stream():
        try:
            while True:
                try:
                    event = await asyncio.wait_for(session.queue.get(), timeout=60.0)
                    session.touch()
                    yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"

                    if event.get("event") in ("complete", "error"):
                        break
                except asyncio.TimeoutError:
                    # 心跳同时刷新 last_active，避免本地大模型慢推理期间
                    # 会话被 sweeper 误清理
                    session.touch()
                    yield f"data: {json.dumps({'event': 'heartbeat'})}\n\n"
        except asyncio.CancelledError:
            pass

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.get("/api/multi-agent/conversation/state/{session_id}")
async def multi_agent_conversation_state(session_id: str):
    """查询多 Agent 协作模式的状态"""
    if session_id not in _multi_agent_sessions:
        raise HTTPException(status_code=404, detail="会话不存在或已过期")

    session = _multi_agent_sessions[session_id]
    if not multi_agent_graph:
        raise HTTPException(status_code=503, detail="Multi-Agent 服务未就绪")

    current = await multi_agent_graph.aget_state(session.config)
    state_values = current.values if current else {}

    chapter_results = state_values.get("chapter_results", []) or []
    return {
        "session_id": session_id,
        "filename": session.filename,
        "current_stage": state_values.get("current_stage", ""),
        "section_count": len(state_values.get("sections", []) or []),
        "total_chapters": state_values.get("total_chapters", 0),
        "total_sections": state_values.get("total_sections", 0),
        "audited_chapters": len(chapter_results),
        "audited_sections": sum(
            len(c.get("subsection_results", []) or []) for c in chapter_results
        ),
        "total_agent_steps": state_values.get("total_agent_steps", 0),
        "finished": state_values.get("finished", False),
        "has_final_report": bool(state_values.get("final_report", "")),
    }


@app.get("/api/multi-agent/conversation/report/{session_id}")
async def multi_agent_conversation_report(session_id: str):
    """获取多 Agent 协作模式的最终报告"""
    if session_id not in _multi_agent_sessions:
        raise HTTPException(status_code=404, detail="会话不存在或已过期")

    session = _multi_agent_sessions[session_id]
    if not multi_agent_graph:
        raise HTTPException(status_code=503, detail="Multi-Agent 服务未就绪")

    current = await multi_agent_graph.aget_state(session.config)
    state_values = current.values if current else {}

    report = state_values.get("final_report", "") or ""
    if not report:
        raise HTTPException(status_code=404, detail="报告尚未生成")

    return {
        "session_id": session_id,
        "report": report,
        "final_report": report,
        "finished": state_values.get("finished", False),
        "chapter_count": len(state_values.get("chapter_results", []) or []),
        "section_count": len(state_values.get("subsection_results", []) or []),
    }


# ============== 逐段审核 API ==============
@app.post("/api/analyze-segment")
async def analyze_segment(
    file: Optional[UploadFile] = File(None),
    session_id: str = Form("default"),
    audit_type: str = Form("risk_management"),
    doc_type: str = Form(""),
    segment_size: str = Form("4000"),  # 接收为字符串，手动转换
    action: str = Form("start"),  # "start"（首次上传）或 "continue"（继续下一段）
    segment_size_q: Optional[int] = Query(None, alias="segment_size", description="每段字符数（Query参数，优先于Form参数）")
):
    """
    逐段审核文档：每次只审核文档的一个段落，支持断点续审

    - action="start": 上传文件并审核第一段
    - action="continue": 继续审核下一段

    Returns:
        段落审核结果 + 进度信息
    """
    # Query 参数优先，否则从 Form 参数手动转换
    if segment_size_q is not None:
        seg_size_val = segment_size_q
    else:
        try:
            seg_size_val = int(segment_size)
        except (ValueError, TypeError):
            seg_size_val = 4000

    if not rag_retriever:
        raise HTTPException(status_code=503, detail="向量库未加载，请稍后重试")

    VALID_AUDIT_TYPES = [
        "risk_management", "design_dev", "software_compliance",
        "registration", "production_quality", "system_construction", "general"
    ]
    if audit_type not in VALID_AUDIT_TYPES:
        audit_type = "general"

    if seg_size_val < 1000:
        seg_size_val = 1000
    if seg_size_val > 8000:
        seg_size_val = 8000

    # 获取审核类型中文标签
    doc_type_label = ""
    if doc_type and rag_retriever:
        doc_type_label = rag_retriever._DOC_TYPE_LABELS.get(doc_type, doc_type)

    if action == "start":
        # ===== 首次上传：提取文本、创建状态、审核第一段 =====
        if not file:
            raise HTTPException(status_code=400, detail="action=start 需要上传文件")

        filename = file.filename or ""
        ext = Path(filename).suffix.lower()
        if ext not in ['.docx', '.pdf']:
            raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

        # 保存并提取文本
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                while True:
                    chunk = await file.read(64 * 1024)
                    if not chunk:
                        break
                    tmp.write(chunk)
                tmp_path = tmp.name
        except Exception:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
            raise

        try:
            from doc_processor import extract_text
            text = extract_text(tmp_path)
            if not text or len(text.strip()) < 50:
                raise HTTPException(status_code=400, detail="文档内容过少或无法提取文本")

            # 创建逐段审核状态
            state = segment_manager.create(
                session_id=session_id,
                document_text=text,
                filename=filename,
                audit_type=audit_type,
                doc_type=doc_type,
                segment_size=seg_size_val
            )

            # 获取第一段
            seg = state.get_next_segment()
            if not seg:
                raise HTTPException(status_code=400, detail="文档内容为空")
            segment_text, seg_idx, total_segs, start_pos, end_pos = seg

            # 审核第一段
            result = await rag_retriever.audit_segment(
                segment_text=segment_text,
                segment_index=seg_idx,
                total_segments=total_segs,
                start_pos=start_pos,
                end_pos=end_pos,
                audit_type=audit_type,
                doc_type_label=doc_type_label
            )
            state.add_result(result)

            progress = state.get_progress()

            # 构建检索文档信息
            retrieved_docs_info = _build_retrieved_info(result.get("relevant_docs", []))

            return {
                "filename": filename,
                "segment_answer": result["answer"],
                "segment_index": seg_idx,
                "total_segments": total_segs,
                "start_pos": start_pos,
                "end_pos": end_pos,
                "progress": progress,
                "action": "start",
                "retrieved_docs": retrieved_docs_info,
                "audit_type": audit_type,
                "full_text_length": len(text)
            }

        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    elif action == "continue":
        # ===== 继续审核：获取下一段 =====
        state = segment_manager.get(session_id)
        if not state:
            raise HTTPException(status_code=404, detail="未找到审核状态，请先上传文件开始审核")

        seg = state.get_next_segment()
        if not seg:
            return {
                "segment_answer": "",
                "segment_index": 0,
                "total_segments": state.total_segments,
                "progress": state.get_progress(),
                "action": "complete",
                "message": "文档已全部审核完成，可生成综合报告"
            }

        segment_text, seg_idx, total_segs, start_pos, end_pos = seg

        result = await rag_retriever.audit_segment(
            segment_text=segment_text,
            segment_index=seg_idx,
            total_segments=total_segs,
            start_pos=start_pos,
            end_pos=end_pos,
            audit_type=state.audit_type,
            doc_type_label=doc_type_label
        )
        state.add_result(result)

        progress = state.get_progress()
        retrieved_docs_info = _build_retrieved_info(result.get("relevant_docs", []))

        return {
            "filename": state.filename,
            "segment_answer": result["answer"],
            "segment_index": seg_idx,
            "total_segments": total_segs,
            "start_pos": start_pos,
            "end_pos": end_pos,
            "progress": progress,
            "action": "continue",
            "retrieved_docs": retrieved_docs_info,
            "audit_type": state.audit_type
        }

    else:
        raise HTTPException(status_code=400, detail=f"不支持的操作: {action}，仅支持 start 或 continue")


def _build_retrieved_info(relevant_docs: List[Dict]) -> List[Dict]:
    """从检索结果构建简要的文档信息列表"""
    seen_sources = set()
    info = []
    for doc in relevant_docs:
        source = doc.get("source", "未知")
        if source not in seen_sources:
            seen_sources.add(source)
            info.append({
                "source": source,
                "preview": doc.get("text", "")[:200]
            })
    return info


@app.get("/api/segment-status/{session_id}")
async def segment_status(session_id: str):
    """查询逐段审核进度"""
    state = segment_manager.get(session_id)
    if not state:
        return {"exists": False, "message": "未找到审核状态"}
    return {
        "exists": True,
        "filename": state.filename,
        "audit_type": state.audit_type,
        "progress": state.get_progress()
    }


@app.post("/api/segment-synthesize")
async def segment_synthesize(
    session_id: str = Form("default")
):
    """
    生成逐段审核的综合报告

    将所有段落审核结果汇总，调用 LLM 生成一份完整的综合审核报告
    """
    if not rag_retriever:
        raise HTTPException(status_code=503, detail="向量库未加载，请稍后重试")

    state = segment_manager.get(session_id)
    if not state:
        raise HTTPException(status_code=404, detail="未找到审核状态，请先上传文件开始审核")

    if not state.segment_results:
        raise HTTPException(status_code=400, detail="还没有审核结果，请先完成至少一段审核")

    # 构建综合上下文
    doc_type_label = ""
    if state.doc_type and rag_retriever:
        doc_type_label = rag_retriever._DOC_TYPE_LABELS.get(state.doc_type, state.doc_type)

    parts = [
        f"# 文档逐段审核综合报告",
        f"**文件名**: {state.filename}",
        f"**审核类型**: {state.audit_type}",
        f"**文件类型**: {doc_type_label}",
        f"**文档总长度**: {len(state.document_text)} 字符",
        f"**审核段落数**: {len(state.segment_results)}",
        "",
        "---",
        "",
        "以下是各段落的审核结果摘要：",
        ""
    ]

    for result in state.segment_results:
        seg_idx = result.get("segment_index", "?")
        answer = result.get("answer", "")
        # 每段截取前1500字符作为综合输入
        summary = answer[:1500] if answer else "(无审核结果)"
        parts.append(f"## 第 {seg_idx} 段审核结果（摘要）")
        parts.append(summary)
        parts.append("")

    parts.append("---")
    parts.append("请基于以上各段审核结果，生成一份完整的综合审核报告。")
    parts.append("报告要求：")
    parts.append("1. 综合概述：概述整个文档的内容和审核范围")
    parts.append("2. 关键发现汇总：汇总各段中发现的关键问题（按严重度排序）")
    parts.append("3. 量化评分：对文档整体进行5维度评分（完整性/规范性/可追溯性/一致性/可操作性）")
    parts.append("4. 修改建议优先级：按P0/P1/P2给出修改优先级建议")
    parts.append("5. 关联法规条款索引")

    synthesis_context = '\n'.join(parts)

    synthesis_prompt = f"""你是一个专业的贴敷式胰岛素泵生产企业文档审核专家。

请根据以下各段落审核结果，生成一份完整的综合审核报告。

注意：
- 综合各段发现，识别跨段落的共性问题
- 对严重问题进行汇总和优先级排序
- 给出文档整体的合规性评价
- 报告应采用Markdown格式，结构清晰"""

    try:
        synthesis_answer = await rag_retriever._call_llm(
            system_prompt=synthesis_prompt,
            user_content=synthesis_context,
            max_tokens=12000,
            temperature=0.5
        )
    except Exception as e:
        logger.error(f"综合报告生成失败: {e}")
        # 降级：手动拼接所有段落结果
        fallback_parts = [f"# 审核报告 — {state.filename}", ""]
        for result in state.segment_results:
            fallback_parts.append(f"## 第 {result.get('segment_index', '?')} 段")
            fallback_parts.append(result.get("answer", ""))
            fallback_parts.append("")
        synthesis_answer = '\n'.join(fallback_parts)

    # 追加各段详细结果
    full_report = synthesis_answer + "\n\n---\n\n# 各段落详细审核结果\n\n"
    for result in state.segment_results:
        seg_idx = result.get("segment_index", "?")
        full_report += f"## 第 {seg_idx} 段\n\n"
        answer = result.get("answer", "")
        # 单段结果截断至3000字符
        full_report += (answer[:3000] if len(answer) > 3000 else answer)
        full_report += "\n\n---\n\n"

    # 限制总报告大小
    MAX_REPORT = 50000
    if len(full_report) > MAX_REPORT:
        full_report = full_report[:MAX_REPORT] + "\n\n⚠️ 报告过长已截断"

    return {
        "synthesis_answer": synthesis_answer,
        "full_report": full_report,
        "segment_count": len(state.segment_results),
        "filename": state.filename,
        "audit_type": state.audit_type
    }


@app.post("/api/segment-reset")
async def segment_reset(session_id: str = Form("default")):
    """重置逐段审核状态"""
    segment_manager.remove(session_id)
    return {"message": "审核状态已重置", "session_id": session_id}


# ============== Markdown 转 DOCX 辅助函数 ==============
def _md_to_docx(md_content: str, doc_title: str) -> "Document":
    """将 Markdown 格式的审核报告转换为 Word 文档"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题页
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(doc_title)
    title_run.bold = True
    title_run.size = Pt(18)
    title_run.font.name = '微软雅黑'
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('医疗器械体系文件审核报告')
    subtitle_run.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    subtitle_run.font.name = '微软雅黑'
    subtitle_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()  # 空行

    # 逐行解析 Markdown
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            heading = doc.add_heading(text, level=min(level, 3))
            for run in heading.runs:
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        # 水平线
        if re.match(r'^[-–—]{3,}\s*$', line.strip()):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            pPr = para._p.get_or_add_pPr()
            pBdr = etree.SubElement(pPr, qn('w:pBdr'))
            bottom = etree.SubElement(pBdr, qn('w:bottom'))
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            i += 1
            continue

        # 无序列表
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if bullet_match:
            para = doc.add_paragraph(style='List Bullet')
            text = bullet_match.group(2)
            _add_formatted_text(para, text)
            i += 1
            continue

        # 有序列表
        numbered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if numbered_match:
            para = doc.add_paragraph(style='List Number')
            text = numbered_match.group(2)
            _add_formatted_text(para, text)
            i += 1
            continue

        # 代码块
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            if code_lines:
                code_para = doc.add_paragraph()
                code_para.paragraph_format.left_indent = Cm(1)
                code_run = code_para.add_run('\n'.join(code_lines))
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # 普通段落
        para = doc.add_paragraph()
        _add_formatted_text(para, line)
        i += 1

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    return doc


def _add_formatted_text(paragraph, text: str):
    """向段落添加带格式的文本（支持 **粗体** 和 `行内代码`）"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    # 分割粗体和普通文本
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
        else:
            run = paragraph.add_run(part)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


# ============== 导出接口 ==============
class ExportRequest(BaseModel):
    """导出请求模型"""
    content: str = Field(..., description="审核报告 Markdown 内容")
    filename: str = Field(default="审核报告", description="原始文件名")

@app.post("/api/export-review")
async def export_review(request: ExportRequest):
    """
    将审核结果导出为 Word (.docx) 文件

    Args:
        request: 包含审核报告 markdown 内容和文件名

    Returns:
        .docx 文件下载
    """
    if not request.content:
        raise HTTPException(status_code=400, detail="审核报告内容为空")

    # 生成报告标题
    base_name = Path(request.filename).stem if request.filename else "审核报告"
    doc_title = f'{base_name} — 审核报告'

    # 转换 Markdown 为 DOCX
    doc = _md_to_docx(request.content, doc_title)

    # 写入内存流
    from io import BytesIO
    from fastapi.responses import Response
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    # 纯 ASCII 文件名
    ascii_name = re.sub(r'[^\x00-\x7F]', '_', base_name).strip('_') or 'report'
    download_name = f'{ascii_name}_audit_report.docx'

    return Response(
        content=buf.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{download_name}"'}
    )


@app.post("/api/clear")
async def clear_history(session_id: str = "default"):
    """清除会话历史"""
    conversation_manager.clear(session_id)
    return {"message": "会话历史已清除", "session_id": session_id}


@app.get("/api/history/{session_id}")
async def get_history(session_id: str = "default"):
    """获取会话历史"""
    history = conversation_manager.get_or_create(session_id)
    return {"session_id": session_id, "history": history}


@app.get("/api/vectorstore/status")
async def vectorstore_status():
    """获取向量库状态"""
    if not vector_store:
        return {"loaded": False, "document_count": 0}

    return {
        "loaded": True,
        "document_count": vector_store.count()
    }


# ============== 运行入口 ==============
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8000,
        workers=1,               # 强制单worker，避免多进程各加载一份HNSW索引
        limit_concurrency=20,    # 限制并发连接数（浏览器需多连接加载页面资源）
        timeout_keep_alive=10,   # 缩短keep-alive，减少空闲连接内存占用
    )
